#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed May 26 15:58:04 2021

@author: ziyuanguan

Usage :  Re-origanize the products folders
         Fetch product properties
         Pack Folders into ZIP waiting for upload.
"""

import pandas as pd 
import subprocess
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from PIL import Image
from docx.enum.text import WD_ALIGN_PARAGRAPH




productColumns = ['产品名称', '产品图片','商城价','库存量','编号','类型','材质','颜色','性别','价格','重量','镜片宽','镜总宽','镜片高','鼻间距','镜腿长','库存量','详情介绍', '分类名称', '标签' ]


defaultPrices = {
    
    '1.56' : '66',
    '1.60' : '90',
    '1.67' : '160',
    '1.74' : '450',
    '非球面': '0',
    '防蓝光' : '50',
    '变灰色': '150',
    '变茶色': '200', 
    '防蓝光变灰色': '200',
    '防蓝光变茶色': '250',
    '偏光灰' : '300',
    '偏光绿' : '300',
    '偏光茶' : '300',
    'other':' 175'
    }

srcFile = '/Users/ziyuanguan/易看-eyecare/零售销售计划/记账.xlsx'
sheetNames = ['鹏光','盛悦']    
sourceDefaultColumns = ['编号','颜色','材质','进货价格','性别','名称','重量','镜片宽','镜总宽','镜片高','鼻间距','镜腿长'] #'进货数量'


inputImagesBase = '../../产品/正式/'
detailImage = ''


date = datetime.strftime(datetime.now() , '%Y-%m-%d')
outputBase = '../../Processed/' + date + '/'
outputImageFolder = outputBase + 'pics/'
outputDetailsFolder = outputBase + 'details/'
outputFolders = [outputImageFolder ,  outputDetailsFolder]

sizeLimit = 1.2 # GB
emphsizedCodes = ['8303']

failed_proc = []

    


def getSize(srcFolder):
    size = int(subprocess.check_output(['du','-sh', srcFolder]).split()[0].decode('utf-8').split('K')[0])
    sizeInG = round(size / 1024 / 1024 , 6)
    return sizeInG
    
    
def zipFiles(srcFolder,name):
    ## example zip -r hello.zip products/*
    command = 'zip -r %s.zip %s*' %(name , srcFolder)
    print('  Zip Command : ' , command)
    proc = subprocess.Popen(command , shell = True)
    if(not proc.returncode == 0):
        failed_proc.append('zip : %s' %srcFolder)


def judgeDetailImage(fileName : str):
    ## 详情或参数图
    if('详情' in fileName or '参数' in fileName):
        return True
    
    
    ## _3 图
    if(len(fileName.split('_')) > 1):
        mid = fileName.split('_')[1]
        
        ## special case
        mid = mid.split('.')[0]
        mid = int(mid)
        if(mid > 30):
            return True
    
    return False


def cleanCache(path , code):
    
    command = 'rm -rf %s' %(outputImageFolder + '*')
    proc = subprocess.Popen(command , shell = True)
    command = 'rm -rf %s' %(outputDetailsFolder + '*')
    proc = subprocess.Popen(command , shell = True)
    command = 'rm -rf %s' %(path + '/' + code)
    proc = subprocess.Popen(command , shell = True)
    command = 'rm -rf %s' %(path + '/detail')
    proc = subprocess.Popen(command , shell = True)
    

def imagesOrderReplace(_code : str , filePath : str , flag:bool = True):
    
    _temp = []
    
    order = [2 , 1]
    
    if(flag):
        ## 正面主图 侧面主图 分别为 _1 _2, 否则反之
        for i in range(len(order)):
            order[i] = 3 - order[i]
    ## 正面主图
    for file in os.listdir(filePath):
        if(_code + '_' + str(order[0]) in file) :
            _temp.append(file)
            
    ## 侧面主图
    for file in os.listdir(filePath):
        if(_code + '_' + str(order[1]) in file) :
            _temp.append(file)
    
    ## 模特图
            
    for file in os.listdir(filePath):
        if('模特' in file) :
            _temp.append(file)
    
    ## 其他买家秀
    for file in os.listdir(filePath):
        if(not file in _temp):
            _temp.append(file)
            
    ## 重命名带有 code
    ## Rename to change order 
    for i in range(len(_temp)) : 
        subprocess.Popen('mv %s %s' %(filePath + '/' + _temp[i] , filePath + '/' + _code + '_' + str(i+1) + '.' + _temp[i].split('.')[-1] ) , shell = True)
    # print('    ---- Image Reorder Successfully')

def detailImageOrderReplace(_code : str , filePath : str):
    _temp = []
    
    ## 参数
    for file in os.listdir(filePath):
        if('参数' in file):
            _temp.append(file)
            
    ## _3x
    
    for file in os.listdir(filePath):
        if(_code + '_3' in file):
            _temp.append(file)

    ## 详情页
    for file in os.listdir(filePath):
        if('详情' in file):
            _temp.append(file)
            
        
    ## 重命名带有 code
    ## Rename to change order 
    for i in range(len(_temp)) : 
        subprocess.Popen('mv %s %s' %(filePath + '/' + _temp[i] , filePath + '/' + _code + '_detail_' + str(i+1) + '.' + _temp[i].split('.')[-1] ) , shell = True)
            
            

def copyImageFolderToDes(srcPath):
    ## example : cp -r ../8301/8301 /Processed/pics/
    
    ## firstly split images into two parts : part1 : which used for main images
    ##                                       part2 : which write into detail docx
    
    ## create two subfolders
    _code = srcPath.split('/')[-1]
    _detailImageFolder = srcPath + '/detail'
    _uploadImageFolder = srcPath + '/' + _code
    _paths = [_detailImageFolder , _uploadImageFolder]

    for path in _paths:
        if(not os.path.exists(path)):
            os.makedirs(path)


    ## split images
    for dirpath, dirnames,filenames in os.walk(srcPath):
        for file in filenames:
            if(file.lower().endswith(('.jpg','.jpeg','.png'))):
                if(judgeDetailImage(file)):
                    subprocess.Popen('cp %s %s' %(srcPath + '/' +file , _detailImageFolder) , shell = True)
                else:
                    subprocess.Popen('cp %s %s' %(srcPath + '/' +file , _uploadImageFolder) , shell = True)
                    
                    
    ## Order the images in main image folder
    imagesOrderReplace(_code , _uploadImageFolder)
    
    ## Order the images in detail folder
    detailImageOrderReplace(_code , _detailImageFolder)
    
    
    command  = 'cp -r %s %s'%(_uploadImageFolder , outputImageFolder)
    proc = subprocess.Popen(command , shell = True)
        
        
def createDetailFile(code : str , srcPath: str):
    ## example : touch /Processed/details/52012.docx
    
    ## put detail image into docx and saved as file
    doc = Document()
    _detailImageFolder = srcPath + '/detail'
    for file in os.listdir(_detailImageFolder):
        #print(_detailImageFolder + '/' + file)
        
        ## save the image into docx
        try:
            doc.add_picture(_detailImageFolder + '/' + file , width = Pt(300) )
        except:
            _jpg = Image.open(_detailImageFolder + '/' + file)
            _jpg.save(_detailImageFolder + '/' + file)
            doc.add_picture(_detailImageFolder + '/' + file , width = Pt(300))
            
        ## align center
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    doc.save(outputDetailsFolder + code + '.docx' )   
    
        
def joinValues(df):
    df = df.astype('str')
    _tempDat = set(df)
    return '、'.join(_tempDat)


def getTypes(row):
    
    materialSrc = row['材质']
    sexSrc = row['性别']
    defaultWord = '、'
    _materials = materialSrc.split('+') 
    _materials = [x + '镜架' for x in _materials]
    _matypes = defaultWord.join(_materials)
    
    if(not sexSrc == 'None'):
        _sexs = sexSrc.split('+')
        _sexs = [x + '士镜架' for x in _sexs]
        _sextypes = defaultWord.join(_sexs)
        
        _matypes += '、'
        _matypes += _sextypes
        
    return _matypes+'、近视镜架'
  

def convertToBetterDigital(src : str):
    
    lastDigit = src[-1]
    if(lastDigit < '3'):
        return src[:-1] + '0'
    
    if(lastDigit >= '3' and lastDigit < '6'):
        return src[:-1] + '6'
    
    if(lastDigit > '6'):
        return src[:-1] + '8'
    
    return src

def generatePriceList( _code : str, price: str):
    
    price = int(price)
    ##强推款 再涨价 20%
    if(_code in emphsizedCodes):
        price = int(1.2 * price)
    _prices = {}
    _indexs = ['1.56','1.60','1.67','1.74']
    
    _services = ['非球面','防蓝光','变灰色','变茶色','防蓝光变灰色','防蓝光变茶色','偏光灰','偏光绿','偏光茶']
    _prices['编号'] = _code
    for index in _indexs:
        for service in _services:
            if(index == '1.74' and '偏光' in service ):
                continue
            else:
                _prices[index + service] = convertToBetterDigital(str(price + int(defaultPrices[index]) + int(defaultPrices[service])))
    
    _prices = pd.DataFrame(_prices, index = [0] )
    
    # print(_prices)
    return _prices
    

def generateProductName(row, flag:bool= True):
    
    
    sex = row['性别']
    name = row['名称']
    code = row['编号']
    ## 中文符号
    sexes = sex.split('+')
    _sex = ''
    if(len(sexes) > 1):
        _sex = '男女通用'
    else:
        _sex = sex + '式'
    
    ## if "盛悦" no need to process
    if(flag):
        return name
    return name + ' ' + _sex + ' ' + code


def processSex(src : str ):
    if(src == '女'):
        return '女式'
    if(src == '男'):
        return '男式'
    return '男女通用'
  
        
def readExcel(srcPath , sheetName):
    data = pd.read_excel(srcPath , sheet_name = sheetName, usecols=sourceDefaultColumns , dtype=str).fillna(method= 'pad')
    data['编号'] = data['编号'].astype('str')
    ## 物流 100， 75 直接加入 镜架价格
    data['进货价格'] = data['进货价格'].map(lambda x : str(int(3.33 * int(x)) + 75))
    new_data = data.groupby(['编号'],as_index=False).agg(joinValues)
    new_data['产品名称'] = new_data.apply(lambda x : generateProductName(x) , axis = 1)
    new_data['产品图片'] = new_data['编号'].map(lambda x : str(x))
    new_data['详情介绍'] = new_data['编号'].map(lambda x : str(x) + '.docx')
    new_data['分类名称'] = new_data.apply(lambda x : getTypes(x) , axis = 1)
    new_data['商城价'] = new_data['进货价格']
    #new_data['规格'] = ''
    #new_data['品牌'] = sheetName
    new_data['型号'] = new_data['编号']
    new_data['库存量'] = 200
    new_data['类型'] = '近视镜架'
    new_data['标签'] = '最新'
    new_data['价格'] = new_data['进货价格']
    new_data['性别'] = new_data['性别'].map(lambda x : processSex(x))
    new_data.rename(columns = {'进货价格':'市场价'}, inplace=True)
    
    ## reorganize the columns orders
    new_data = new_data[productColumns]
    
    return new_data
    

    

## get product details from '进货单'
## copy image and detail folders
## generate list.xlsx
## split size if too large
## packed as ZIP
def run(srcPath , sheetName , clean : bool = 'False'):
    
    ## create folders if not prepared
    for folder in outputFolders :
        if(not os.path.exists(folder)):
            os.makedirs(folder)
    
    productsInfo = readExcel(srcPath , sheetName)
    print(productsInfo)
    outputListPath = outputBase + 'list.xlsx'
    productsInfo.to_excel(outputListPath, header = True , index = False)
    
    ## copy & create files 
    _codes = productsInfo['编号'].tolist()
    _priceList = pd.DataFrame([])
    for _code in _codes :
        
        ##cleanCache
        if(clean):
            cleanCache(path=inputImagesBase + sheetName + '/' + _code ,  code = _code)
        else:
            # copy images
            copyImageFolderToDes(inputImagesBase + sheetName + '/' + _code)
            ## create detail file
            createDetailFile(_code , inputImagesBase + sheetName + '/' + _code)
            # generate price list 
            _tempPrice = generatePriceList(_code, productsInfo.loc[productsInfo['编号'] == _code]['价格'].values[0])
            _priceList = _priceList.append(_tempPrice)
            
    _priceList.to_csv('../../Processed/' + 'priceList' + sheetName + '.csv', index=False)
    ## zip files
    # zipFiles(outputBase , date)
               
    print('fail processes : ', failed_proc)
    # print(_priceList)
            
run(srcFile,'盛悦', False)